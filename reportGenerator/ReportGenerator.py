#!/usr/bin/python3.6
# -*- coding: utf-8 -*-

import os
import datetime
import csv
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
document = Document()

#Class that contains scan parameters.
class Scan:
	def __init__(self, title ='', message ='', message2 =''):
		self.title = title
		self.message = message
		self.message2 = message2

def initPerPage(title, message, message2):
	document.add_page_break()
	document.add_heading(title, 1)
	document.add_paragraph('')
	document.add_paragraph(message)
	document.add_paragraph(message2)
	document.add_paragraph('')

def createPageTestSSL(title, message, message2, ipDomain):
	initPerPage(title, message, message2)
	pathAndFileName = "/temp/testssl_f_" + ipDomain + ".json"

	try:
		f = open(pathAndFileName)
		s = f.read()
		jsonResults = json.loads(s)
		f.close()
	except:
		p = document.add_paragraph('')
		p.add_run('There is not results for the TestSSL scanner.').bold = True
	else:
		#Parsing vulnerabilities here
		table = document.add_table(rows=1, cols=4)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'ID'
		hdr_cells[1].text = 'Port'
		hdr_cells[2].text = 'Severity'
		hdr_cells[3].text = 'Finding'

		for num in range(len(jsonResults)):
			row_cells = table.add_row().cells
			row_cells[0].text = jsonResults[num]['id']
			row_cells[1].text = jsonResults[num]['port']
			row_cells[2].text = jsonResults[num]['severity']
			row_cells[3].text = jsonResults[num]['finding']


def createPageSSHAudit(title, message, message2, ipDomain):
	initPerPage(title, message, message2)
	pathAndFileName = "/temp/sshaudit_" + ipDomain + ".txt"

	try:
		f = open(pathAndFileName)
		data = f.readlines()
		f.close()
		if "[exception]" not in data[0]:
			document.add_paragraph(data)
		else:
			p = document.add_paragraph('')
			p.add_run('There is not SSH service available.').bold = True
	except:
		p = document.add_paragraph('')
		p.add_run('There is not SSH service available.').bold = True


def createPageWhatWeb(title, message, message2, ipDomain, level):
	initPerPage(title, message, message2)
	pathAndFileName = "/temp/whatweb_f_" + level + "_" + ipDomain + ".json"

	try:
		f = open(pathAndFileName)
		s = f.read()
		jsonResults = json.loads(s)
		f.close()
	except:
		p = document.add_paragraph('')
		p.add_run('There is not WhatWeb results.').bold = True
	else:
		#Parsing vulnerabilities here
		table = document.add_table(rows=1, cols=2)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'Type'
		hdr_cells[1].text = 'Information'
		
		#Loop to create the table with the JSON results
		for key, value in jsonResults.items():
			#print("key: {}, value: {}".format(key, value))
			row_cells = table.add_row().cells
			row_cells[0].text = key
			for key2, value2 in value.items():
				#print("key2: {}, value2: {}".format(key2, value2))
				row_cells[1].text += '\n' + key2 + ": " + ''.join(value2)

def createPageDroopescan(title, message, message2, ipDomain):
	initPerPage(title, message, message2)
	pathAndFileName = "/temp/droopscan_f_" + ipDomain + ".json"

	try:
		f = open(pathAndFileName)
		data = f.readlines()
		f.close()
		if "not identified CMS" not in data[0]:
			document.add_paragraph(data)
		else:
			p = document.add_paragraph('')
			p.add_run('There is not identified CMS in the selected target.').bold = True
	except:
		p = document.add_paragraph('')
		p.add_run('There are not CMS.').bold = True

def createPageNikto(title, message, message2, ipDomain, level):
	initPerPage(title, message, message2)
	pathAndFileName = "/temp/niktoscan_f_" + level + "_" + ipDomain + ".json"

	try:
		f = open(pathAndFileName)
		s = f.read()
		jsonResults = json.loads(s)
		f.close()
	except:
		p = document.add_paragraph('')
		p.add_run('There is not results for the Nikto scanner.').bold = True
	else:
		#Parsing vulnerabilities here
		table = document.add_table(rows=1, cols=5)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'ID'
		hdr_cells[1].text = 'Title'
		hdr_cells[2].text = 'Port'
		hdr_cells[3].text = 'Petition Type'
		hdr_cells[4].text = 'URL'

		for num in range(len(jsonResults)):
			row_cells = table.add_row().cells
			row_cells[0].text = jsonResults[num]['id']
			row_cells[1].text = jsonResults[num]['tittle']
			row_cells[2].text = jsonResults[num]['port']
			row_cells[3].text = jsonResults[num]['petitionType']
			row_cells[4].text = jsonResults[num]['url']

def createPageWPScan(title, message, message2, ipDomain, level):
	initPerPage(title, message, message2)
	pathAndFileName = "/temp/wpscan_f_" + level + "_" + ipDomain + ".json"
	jsonResults = ""

	try:
		f = open(pathAndFileName)
		s = f.read()
		jsonResults = json.loads(s)
		f.close()
		if "does not seem to be running WordPress" in str(jsonResults):
			p = document.add_paragraph('')
			p.add_run(str(jsonResults)).bold = True
		else:
			#Parsing vulnerabilities here
			table = document.add_table(rows=1, cols=2)
			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = 'Type'
			hdr_cells[1].text = 'Information'
			
			#Loop to create the table with the JSON results
			for key, value in jsonResults.items():
				#print("key: {}, value: {}".format(key, value))
				row_cells = table.add_row().cells
				row_cells[0].text = key
				for key2, value2 in value.items():
					#print("key2: {}, value2: {}".format(key2, value2))
					row_cells[1].text += '\n' + key2 + ": " + ''.join(value2)
	except:
		p = document.add_paragraph('')
		p.add_run('There is not WPScan results.').bold = True

def createPageOpenVAS(title, message, message2, ipDomain, level):
	initPerPage(title, message, message2)
	pathAndFileName = "/temp/openvas_f_" + level + "_" + ipDomain + ".json"

	try:
		f = open(pathAndFileName)
		s = f.read()
		jsonResults = json.loads(s)
		f.close()
	except:
		p = document.add_paragraph('')
		p.add_run('There is not results for the OpenVAS scanner.').bold = True
	else:
		for num in range(len(jsonResults)):
			#Parsing vulnerabilities here
			table = document.add_table(rows=1, cols=2)
			row_cells = table.rows[0].cells
			row_cells[0].text = 'CVE'
			row_cells[1].text = jsonResults[num]['cve']
			row_cells = table.add_row().cells
			row_cells[0].text = 'Title'
			row_cells[1].text = jsonResults[num]['tittle']
			row_cells = table.add_row().cells
			row_cells[0].text = 'Port'
			row_cells[1].text = jsonResults[num]['port']
			row_cells = table.add_row().cells
			row_cells[0].text = 'Protocol'
			row_cells[1].text = jsonResults[num]['protocol']
			row_cells = table.add_row().cells
			row_cells[0].text = 'Severity'
			row_cells[1].text = jsonResults[num]['severity']
			row_cells = table.add_row().cells
			row_cells[0].text = 'CVSS'
			row_cells[1].text = jsonResults[num]['cvss']
			row_cells = table.add_row().cells
			row_cells[0].text = 'Description'
			row_cells[1].text = jsonResults[num]['description']
			row_cells = table.add_row().cells
			row_cells[0].text = 'Impact'
			row_cells[1].text = jsonResults[num]['impact']
			row_cells = table.add_row().cells
			row_cells[0].text = 'Solution'
			row_cells[1].text = jsonResults[num]['solution']
			document.add_paragraph('')

def coverPage(ipDomain, level):
	timeNow = datetime.datetime.now()

	for i in range(10):
		document.add_paragraph('')

	title = document.add_heading('Aut0test',0)
	subtitle = document.add_paragraph('Report of finding', style='Subtitle')

	for i in range(3):
		document.add_paragraph('')

	document.add_paragraph('Target: ' + ipDomain)
	if level == "1":
		document.add_paragraph('Level of scan: Low')
	if level == "2":
		document.add_paragraph('Level of scan: Medium')
	if level == "3":
		document.add_paragraph('Level of scan: High')

	for i in range(7):
		document.add_paragraph('')
	
	date = document.add_paragraph('Date: ' + timeNow.strftime("%d-%m-%Y %H:%M"))
	date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

#Program start here
def startReport(ipDomain, level):
	scans = []

	#TestSSL object
	scans.append(Scan('TestSSL Results', 'With TestSSL we can check vulnerabilities and information in TLS/SSL communications, service on ports, ciphers, protocols, cryptographic flaws and more.', 'In this section you could check the results of the test.'))
	#SSH-Audit object
	scans.append(Scan('SSH-Audit Results', 'SSH-Audit is an script to search for vulnerabilities and information in SSH1 and SSH2 protocol.', 'In this section you could check the results of the test.'))
	#WhatWeb object
	scans.append(Scan('WhatWeb Results', 'WhatWeb was design to recognises web technologies, CMSs, blogging platforms, statistic/analytics packages, JavaScript libraries, web servers, and embedded devices. Also identifies version numbers, email addresses, account IDs, web framework modules, SQL errors, and more interesting information.', 'In this section you could check the results of the test.'))	
	#Droopescan object
	scans.append(Scan('Droopescan Results', 'Droopescan is a scanner to identify issues with several CMS like: Drupal, SolverStripe and WordPress.', 'In this section you could check the results of the scan.'))
	#Nikto object
	scans.append(Scan('Nikto Results', 'Nikto is a web scanner which checks for outdated/deprecated versions, server configurations, it identifies installed web servers, software, plugins among other functionalities.', 'In this section you could check the results of the scan.'))
	#WPScan object
	scans.append(Scan('WPScan Results', 'WPScan is a tool to perform web vulnerability scans for WordPress websites.', 'In this section you could check the results of the scan.'))
	#OpenVAS object
	scans.append(Scan('OpenVAS Results', 'OpenVAS is a framework of several services and tools to launch web vulnerability scans.', 'In this section you could check the results of the scan.'))

	#Here we print the cover page
	coverPage(ipDomain,level)

	#Here we print the pages for each scan
	createPageTestSSL(scans[0].title, scans[0].message, scans[0].message2, ipDomain)
	createPageSSHAudit(scans[1].title, scans[1].message, scans[1].message2, ipDomain)
	createPageWhatWeb(scans[2].title, scans[2].message, scans[2].message2, ipDomain, level)
	createPageDroopescan(scans[3].title, scans[3].message, scans[3].message2, ipDomain)
	createPageNikto(scans[4].title, scans[4].message, scans[4].message2, ipDomain, level)
	createPageWPScan(scans[5].title, scans[5].message, scans[5].message2, ipDomain, level)
	createPageOpenVAS(scans[6].title, scans[6].message, scans[6].message2, ipDomain, level)

	document.add_page_break()

	newPathAndFileName = "reportAut0test.docx"
	#Delete file if exists
	if os.path.isfile(newPathAndFileName):
		os.remove(newPathAndFileName)

	document.save(newPathAndFileName)
